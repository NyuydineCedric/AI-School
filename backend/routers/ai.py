import os
import uuid
import mimetypes
from pathlib import Path

from dotenv import load_dotenv
from langchain_google_genai import ChatGoogleGenerativeAI
from fastapi import APIRouter, Depends, UploadFile, File, HTTPException
from io import BytesIO
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List

from auth import get_current_user, require_role
import models

load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

router = APIRouter(tags=["ai"])


def extract_text(content) -> str:
    """Gemini sometimes returns content as a list of blocks (e.g.
    [{'type': 'text', 'text': '...'}]) instead of a plain string. Flatten
    that down to plain text either way."""
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        parts = []
        for chunk in content:
            if isinstance(chunk, str):
                parts.append(chunk)
            elif isinstance(chunk, dict) and "text" in chunk:
                parts.append(chunk["text"])
        return "".join(parts)
    return str(content)



def _iter_block_items(document):
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    parent_elm = document.element.body
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _extract_docx_text(doc) -> str:
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    parts: list[str] = []
    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            if block.text.strip():
                parts.append(block.text)
        elif isinstance(block, Table):
            for row in block.rows:
                cells = [c.text.strip() for c in row.cells]
                line = " | ".join(cells)
                if line.strip(" |"):
                    parts.append(line)
    return "\n".join(parts)


class ChatMessage(BaseModel):
    role: str  # "user" or "assistant"
    content: str


class Question(BaseModel):
    messages: List[ChatMessage]
    accent: str = "USA"
    temperature: float = 1.0


@router.post("/ai")
def ask_ai(q: Question, user: models.User = Depends(get_current_user)):
    lim = ChatGoogleGenerativeAI(
        model="gemini-3.1-flash-lite",
        google_api_key=GOOGLE_API_KEY,
        temperature=q.temperature,
    )

    message = [
        ("system", f"You are an expert educational AI assistant and tutor. Help students and teachers with course materials, explanations, questions, and document generation for any subject. Reply in a {q.accent} English tone. When generating content for documents, provide well-structured, clear, and educational material.")
    ]
    for m in q.messages:
        message.append((m.role, m.content))

    def token_stream():
        for part in lim.stream(message):
            if isinstance(part.content, str):
                yield part.content
            elif isinstance(part.content, list):
                for chunk in part.content:
                    if isinstance(chunk, str):
                        yield chunk
                    elif isinstance(chunk, dict) and "text" in chunk:
                        yield chunk["text"]

    return StreamingResponse(token_stream(), media_type="text/plain")


class GenerateNotesRequest(BaseModel):
    course_name: str
    topic: str = ""


class GenerateNotesResponse(BaseModel):
    content: str


@router.post("/ai/generate-notes", response_model=GenerateNotesResponse)
def generate_notes(
    payload: GenerateNotesRequest,
    user: models.User = Depends(require_role("teacher")),
):
    llm = ChatGoogleGenerativeAI(
        model="gemini-3.1-flash-lite",
        google_api_key=GOOGLE_API_KEY,
        temperature=0.6,
    )

    prompt = f'''Write engaging study notes for "{payload.course_name}".

IMPORTANT INSTRUCTIONS:
- Write in a natural, conversational tone as if written by an experienced teacher
- Avoid clichés like "In this lesson" or "It is important to understand"
- Use varied sentence structures and active voice
- Include practical examples when relevant
- Organize content clearly but naturally
- Write roughly 100-300 words with substance
- Focus on genuine understanding, not just definitions
- Avoid sounding robotic or overly formal
- Include real world examples 
- Keep it short and simple
- Answer based on the way the question is asked

{f'TOPIC: {payload.topic.strip()}' if payload.topic.strip() else 'Cover fundamental concepts'}

Make it engaging and educational:'''

    response = llm.invoke([
        ("system", "You are an excellent educator who writes study notes that are engaging, clear, and genuinely help students learn. Write naturally without mentioning you are AI or using disclaimers. Use your expertise to provide real educational value."),
        ("user", prompt),
    ])
    content = extract_text(response.content)
    return GenerateNotesResponse(content=content)


class GenerateDocumentRequest(BaseModel):
    course_name: str
    topic: str = ""
    format: str = "pdf"  # 'pdf' or 'docx'
    title: str | None = None


@router.post("/ai/generate-document")
def generate_document(
    payload: GenerateDocumentRequest,
    user: models.User = Depends(require_role("teacher")),
):
    llm = ChatGoogleGenerativeAI(
        model="gemini-3.1-flash-lite",
        google_api_key=GOOGLE_API_KEY,
        temperature=0.6,
    )

    prompt = f'''Write comprehensive, engaging study notes for "{payload.course_name}".

IMPORTANT INSTRUCTIONS:
- Write in a natural, conversational tone as if written by an experienced educator
- Avoid generic phrases like "In this section" or "It is important to note"
- Use varied sentence structures and avoid repetition
- Include real-world examples and practical applications
- Organize content logically with clear sections
- Use transitions between ideas naturally
- Write roughly 800-1000 words with depth and substance
- Focus on helping students truly understand the material
- Avoid mentioning this is "AI-generated" or using disclaimers

{f'SPECIFIC TOPIC: Focus primarily on {payload.topic.strip()}' if payload.topic.strip() else 'Cover the main concepts and fundamentals of the subject'}

Structure the content with:
1. An engaging introduction that hooks the reader
2. 3-5 main sections with natural section headers (not numbered)
3. Practical examples and applications
4. Key takeaways or conclusion

Write this now:'''

    response = llm.invoke([
        ("system", "You are an experienced educator writing study materials. Your goal is to make content engaging, informative, and easy to understand. Write naturally without sounding robotic or like it was generated by AI. Use your expertise to provide genuine educational value."),
        ("user", prompt),
    ])
    content = extract_text(response.content)

    # DOCX
    if payload.format.lower() == "docx":
        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except Exception:
            raise HTTPException(status_code=500, detail="python-docx not installed on the server")

        doc = DocxDocument()

        # Add title
        title = payload.title or payload.course_name
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata
        if payload.topic.strip():
            meta_para = doc.add_paragraph(f"Topic: {payload.topic.strip()}")
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta_style = meta_para.style
            meta_para.runs[0].font.size = Pt(10)
            meta_para.runs[0].font.color.rgb = RGBColor(100, 100, 100)

        # Add spacing
        doc.add_paragraph()

        # Process content - intelligently parse for sections
        lines = content.split("\n")
        current_section = None

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Detect section headers (lines that are likely headers)
            if (len(line) < 80 and
                line[0].isupper() and
                not line.endswith(".") and
                not any(line.startswith(f"{i}.") for i in range(10))):
                # This looks like a header
                doc.add_heading(line, level=2)
                current_section = line
            elif line.startswith("•") or line.startswith("-"):
                # Bullet point
                bullet_text = line.lstrip("•- ").strip()
                doc.add_paragraph(bullet_text, style='List Bullet')
            else:
                # Regular paragraph
                para = doc.add_paragraph(line)
                # Add some spacing for readability
                para_format = para.paragraph_format
                para_format.space_after = Pt(6)
                para_format.line_spacing = 1.15

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        filename = f"{(payload.title or payload.course_name).replace(' ', '_')}.docx"
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=\"{filename}\""},
        )

    # PDF
    if payload.format.lower() == "pdf":
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.units import inch
            from reportlab.lib.styles import ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
        except Exception:
            raise HTTPException(status_code=500, detail="reportlab not installed on the server")

        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=1*inch,
            bottomMargin=0.75*inch,
        )

        story = []

        # Styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=ParagraphStyle(
                'Normal',
                fontName='Helvetica-Bold',
                fontSize=24,
                textColor='#1F2937',
            ),
            alignment=TA_CENTER,
            spaceAfter=12,
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            fontName='Helvetica-Bold',
            fontSize=13,
            textColor='#1F2937',
            spaceAfter=6,
            spaceBefore=12,
        )

        normal_style = ParagraphStyle(
            'CustomNormal',
            fontName='Helvetica',
            fontSize=10.5,
            alignment=TA_JUSTIFY,
            spaceAfter=8,
            leading=14,
        )

        # Add title
        title = payload.title or payload.course_name
        story.append(Paragraph(title, title_style))

        # Add metadata
        if payload.topic.strip():
            meta_style = ParagraphStyle(
                'Meta',
                fontName='Helvetica',
                fontSize=9,
                textColor='#666666',
                alignment=TA_CENTER,
            )
            story.append(Paragraph(f"<i>Topic: {payload.topic.strip()}</i>", meta_style))

        story.append(Spacer(1, 0.2*inch))

        # Process content
        lines = content.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 0.1*inch))
                continue

            # Detect headers
            if (len(line) < 80 and
                line[0].isupper() and
                not line.endswith(".") and
                not any(line.startswith(f"{i}.") for i in range(10))):
                story.append(Paragraph(line, heading_style))
            elif line.startswith("•") or line.startswith("-"):
                bullet_text = line.lstrip("•- ").strip()
                bullet_style = ParagraphStyle(
                    'BulletStyle',
                    fontName='Helvetica',
                    fontSize=10.5,
                    leftIndent=20,
                    spaceAfter=6,
                    leading=14,
                )
                story.append(Paragraph(f"• {bullet_text}", bullet_style))
            else:
                story.append(Paragraph(line, normal_style))

        doc.build(story)
        buffer.seek(0)
        filename = f"{(payload.title or payload.course_name).replace(' ', '_')}.pdf"
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=\"{filename}\""},
        )

    raise HTTPException(status_code=400, detail="Unsupported format; use 'pdf' or 'docx'")


@router.post("/ai/convert-to-docx")
async def convert_to_docx(file: UploadFile = File(...), user: models.User = Depends(require_role("teacher"))):
    filename = file.filename or "uploaded"
    name = filename.rsplit(".", 1)[0]

    # TXT -> DOCX (simple)
    if filename.lower().endswith(".txt"):
        text = (await file.read()).decode("utf-8", errors="ignore")
        try:
            from docx import Document as DocxDocument
        except Exception:
            raise HTTPException(status_code=500, detail="python-docx not installed on the server")

        doc = DocxDocument()
        for line in text.split("\n"):
            doc.add_paragraph(line)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        out_name = f"{name}.docx"
        return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename=\"{out_name}\""})

    # DOCX -> return back as-is (no conversion needed)
    if filename.lower().endswith(".docx"):
        data = await file.read()
        return StreamingResponse(BytesIO(data), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename=\"{name}_copy.docx\""})

    # PDF -> extract text and produce DOCX
    if filename.lower().endswith(".pdf"):
        try:
            from PyPDF2 import PdfReader
        except Exception:
            raise HTTPException(status_code=500, detail="PyPDF2 not installed on the server")

        raw = await file.read()
        reader = PdfReader(BytesIO(raw))
        pages = []
        for p in reader.pages:
            try:
                pages.append(p.extract_text() or "")
            except Exception:
                pages.append("")
        text = "\n\n".join(pages)

        try:
            from docx import Document as DocxDocument
        except Exception:
            raise HTTPException(status_code=500, detail="python-docx not installed on the server")

        doc = DocxDocument()
        for line in text.split("\n"):
            doc.add_paragraph(line)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        out_name = f"{name}.docx"
        return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename=\"{out_name}\""})

    raise HTTPException(status_code=400, detail="Unsupported file type for conversion")


class ExtractTextResponse(BaseModel):
    content: str


@router.post("/ai/extract-text", response_model=ExtractTextResponse)
async def extract_text_from_upload(
    file: UploadFile = File(...),
    user: models.User = Depends(require_role("teacher")),
):
    """Extract text from uploaded PDF, DOCX, DOC, or TXT files"""
    filename = file.filename or "uploaded"

    try:
        if filename.lower().endswith(".txt"):
            content = (await file.read()).decode("utf-8", errors="ignore")
            return ExtractTextResponse(content=content)

        elif filename.lower().endswith(".docx"):
            try:
                from docx import Document as DocxDocument
            except Exception:
                raise HTTPException(status_code=500, detail="python-docx not installed")

            raw = await file.read()
            doc = DocxDocument(BytesIO(raw))
            content = _extract_docx_text(doc)
            return ExtractTextResponse(content=content)

        elif filename.lower().endswith(".doc"):
            # NOTE: python-docx cannot open legacy binary .doc files at all
            # (it only understands the modern .docx/OOXML zip format). This
            # branch previously tried DocxDocument() on a .doc file, which
            # threw and got swallowed elsewhere, silently returning nothing.
            # Being explicit here instead of pretending it's supported.
            raise HTTPException(
                status_code=400,
                detail="Legacy .doc files aren't supported — please re-save as .docx and re-upload.",
            )

        elif filename.lower().endswith(".pdf"):
            try:
                from PyPDF2 import PdfReader
            except Exception:
                raise HTTPException(status_code=500, detail="PyPDF2 not installed")

            raw = await file.read()
            reader = PdfReader(BytesIO(raw))
            pages = []
            for page in reader.pages:
                try:
                    text = page.extract_text()
                    if text:
                        pages.append(text)
                except Exception:
                    pass

            content = "\n".join(pages)
            return ExtractTextResponse(content=content)

        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Use PDF, DOCX, or TXT")

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to extract text: {str(e)}")


# ---------------------------------------------------------------------------
# AI Tutor chat file upload
# ---------------------------------------------------------------------------
# Accepts images, audio, and common document formats (PDF/DOC/DOCX/PPT/
# PPTX/XLS/XLSX/TXT/CSV) for the AI Tutor chat. Video is intentionally NOT
# accepted here — video files are too large to usefully upload and the
# model can't watch them anyway. The frontend instead lets the user paste
# a video link (e.g. YouTube), which is passed to the model as a plain URL
# reference in the chat message text.
#
# For text-based formats (PDF/DOCX/DOC/TXT/CSV) and images, text is
# extracted (via OCR as a fallback for scans/photos, and directly for
# images) and returned alongside the file metadata so the AI Tutor chat
# endpoint (POST /ai above) actually has real content to work with,
# instead of just a filename. Extraction is capped at MAX_EXTRACTED_CHARS
# to avoid blowing up the model's context window on very long documents.
#
# OCR requires the Tesseract OCR engine installed on the machine (a
# system-level program, not a pip package) — see setup notes alongside
# this file.

UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

MAX_UPLOAD_SIZE_BYTES = 25 * 1024 * 1024  # 25 MB
MAX_EXTRACTED_CHARS = 12000

_CHAT_UPLOAD_ALLOWED_PREFIXES = ("image/", "audio/")
_CHAT_UPLOAD_ALLOWED_EXACT_TYPES = {
    "application/pdf",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.ms-powerpoint",
    "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    "application/vnd.ms-excel",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "text/plain",
    "text/csv",
}


def _chat_upload_type_allowed(content_type: str) -> bool:
    if not content_type:
        return False
    return (
        content_type.startswith(_CHAT_UPLOAD_ALLOWED_PREFIXES)
        or content_type in _CHAT_UPLOAD_ALLOWED_EXACT_TYPES
    )


def _ocr_image_bytes(raw: bytes) -> str | None:
    """Run OCR on a raw image (photo/scan) and return extracted text, or
    None if OCR isn't available or finds nothing."""
    try:
        import pytesseract
        from PIL import Image

        img = Image.open(BytesIO(raw))
        text = pytesseract.image_to_string(img)
        return text.strip() or None
    except Exception:
        return None


def _ocr_pdf_bytes(raw: bytes) -> str | None:
    """Render each page of a PDF to an image and OCR it. Used as a fallback
    when a PDF has no embedded text layer (i.e. it's a scan/photo saved as
    PDF rather than a digitally generated document)."""
    try:
        import fitz  # PyMuPDF
        import pytesseract
        from PIL import Image

        doc = fitz.open(stream=raw, filetype="pdf")
        page_texts = []
        for page in doc:
            pix = page.get_pixmap(dpi=200)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            page_texts.append(pytesseract.image_to_string(img))
        doc.close()
        combined = "\n".join(page_texts).strip()
        return combined or None
    except Exception:
        return None


def _extract_chat_upload_text(filename: str, content_type: str, raw: bytes) -> str | None:
    """Best-effort text extraction so uploaded documents are actually
    readable by the model. Falls back to OCR for scanned PDFs and for
    image uploads (photos of documents). Returns None only if nothing
    could be extracted at all."""
    lower_name = filename.lower()

    try:
        if lower_name.endswith(".txt") or lower_name.endswith(".csv"):
            return raw.decode("utf-8", errors="ignore")

        if lower_name.endswith(".pdf"):
            from PyPDF2 import PdfReader

            reader = PdfReader(BytesIO(raw))
            pages = []
            for page in reader.pages:
                try:
                    text = page.extract_text()
                    if text:
                        pages.append(text)
                except Exception:
                    pass
            direct_text = "\n".join(pages).strip()

            # A real (non-scanned) PDF usually yields a decent amount of
            # text. If we got little/nothing, this is likely a scan or
            # photo saved as PDF — fall back to OCR.
            if len(direct_text) >= 20:
                return direct_text
            return _ocr_pdf_bytes(raw) or (direct_text or None)

        if lower_name.endswith(".docx"):
            from docx import Document as DocxDocument

            doc = DocxDocument(BytesIO(raw))
            content = _extract_docx_text(doc)
            return content or None

        if lower_name.endswith(".doc"):

            return None

        if content_type.startswith("image/"):
           
            return _ocr_image_bytes(raw)

    except Exception:
       
        return None

    return None  

@router.post("/ai/chat/upload")
async def upload_chat_file(
    file: UploadFile = File(...),
    user: models.User = Depends(get_current_user),
):
    content_type = file.content_type or mimetypes.guess_type(file.filename or "")[0] or ""
    if not _chat_upload_type_allowed(content_type):
        raise HTTPException(
            status_code=415,
            detail=(
                f"File type '{content_type}' is not allowed. "
                "For videos, paste a link instead of uploading the file."
            ),
        )

    contents = await file.read()
    size = len(contents)
    if size > MAX_UPLOAD_SIZE_BYTES:
        raise HTTPException(status_code=413, detail="File exceeds the 25MB limit.")

    file_id = str(uuid.uuid4())
    original_name = file.filename or "upload"
    stored_name = f"{file_id}{Path(original_name).suffix}"

    with open(UPLOAD_DIR / stored_name, "wb") as f:
        f.write(contents)

    extracted_text = _extract_chat_upload_text(original_name, content_type, contents)
    truncated = False
    if extracted_text and len(extracted_text) > MAX_EXTRACTED_CHARS:
        extracted_text = extracted_text[:MAX_EXTRACTED_CHARS]
        truncated = True

    return {
        "id": file_id,
        "name": original_name,
        "url": f"/uploads/{stored_name}",
        "type": content_type,
        "size": size,
        "extracted_text": extracted_text,
        "truncated": truncated,
    }